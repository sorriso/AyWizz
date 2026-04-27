# =============================================================================
# File: parser.py
# Version: 2
# Path: ay_platform_core/src/ay_platform_core/c7_memory/ingestion/parser.py
# Description: MIME-type dispatching parsers. v2 (Phase B of v1 plan)
#              activates the production parsers for PDF (pypdf),
#              HTML (BeautifulSoup with lxml when available, html.parser
#              fallback), and DOCX (python-docx). text/plain and
#              text/markdown remain pure-stdlib.
#
#              All parsers consume `bytes` and return UTF-8 string.
#              Failures raise — the caller (ingestion service) maps
#              them to HTTP 415 (UnsupportedMimeError) or 422
#              (ParseFailureError).
#
#              v2 explicitly drops the image OCR path: in v1 functional
#              scope, images are not accepted at the upload boundary;
#              R-400-002 OCR is reserved for v1.5+.
#
# @relation implements:R-400-021
# =============================================================================

from __future__ import annotations

import io
import re
from collections.abc import Callable

# Markdown frontmatter fence (`--- ... ---` at file start).
_FRONTMATTER_FENCE = re.compile(r"^---\n.*?\n---\n?", re.DOTALL)

# DOCX MIME — the official OpenXML wordprocessing type.
_DOCX_MIME = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


class UnsupportedMimeError(RuntimeError):
    """Raised by the dispatcher when no parser is registered for a MIME type."""


class ParseFailureError(RuntimeError):
    """Raised when a registered parser cannot extract text from the bytes
    (corrupt PDF, malformed HTML, encrypted DOCX, etc.)."""


def parse(mime_type: str, content: bytes) -> str:
    """Dispatch to the parser matching `mime_type`.

    Returns the extracted plain text. Any failure raises — the caller
    (ingestion service) maps that to a `failed` source record.
    """
    parser = _REGISTRY.get(mime_type)
    if parser is None:
        raise UnsupportedMimeError(
            f"no parser registered for MIME type {mime_type!r}; "
            f"v1 supports: {', '.join(sorted(_REGISTRY))}"
        )
    return parser(content)


def supported_mime_types() -> tuple[str, ...]:
    """The MIME types accepted at the upload boundary."""
    return tuple(sorted(_REGISTRY))


# ---------------------------------------------------------------------------
# Concrete parsers
# ---------------------------------------------------------------------------


def _parse_text(content: bytes) -> str:
    return content.decode("utf-8", errors="replace")


def _parse_markdown(content: bytes) -> str:
    """Strip the YAML frontmatter if present, keep the rest as-is."""
    text = content.decode("utf-8", errors="replace")
    return _FRONTMATTER_FENCE.sub("", text, count=1).lstrip()


def _parse_pdf(content: bytes) -> str:
    """Extract text from a PDF using `pypdf`. Pure-Python, no native deps.

    Concatenates page text with double-newline separators so chunkers
    can detect page boundaries. Encrypted PDFs raise ParseFailureError —
    we don't support password-protected sources in v1.
    """
    # Local imports keep cold-start fast: pypdf brings several MB of
    # PDF parsing code that we don't want loading every C7 process even
    # when no PDF upload happens.
    from pypdf import PdfReader  # noqa: PLC0415
    from pypdf.errors import PdfReadError  # noqa: PLC0415

    try:
        reader = PdfReader(io.BytesIO(content))
    except PdfReadError as exc:
        raise ParseFailureError(f"invalid PDF: {exc}") from exc
    if reader.is_encrypted:
        raise ParseFailureError(
            "encrypted PDF — password-protected sources are not supported in v1"
        )
    pages: list[str] = []
    for i, page in enumerate(reader.pages):
        try:
            pages.append(page.extract_text() or "")
        except Exception as exc:
            # pypdf can raise a variety of stream-decoding errors on
            # malformed pages; degrade gracefully — record the page as
            # empty rather than failing the whole document.
            pages.append(f"[page {i + 1} unreadable: {type(exc).__name__}]")
    return "\n\n".join(pages).strip()


def _parse_html(content: bytes) -> str:
    """Extract text from HTML using BeautifulSoup. Drops <script>/<style>
    blocks (their text content is irrelevant to retrieval) and collapses
    whitespace. Uses `lxml` parser when available (transitively pulled
    in by python-docx); falls back to stdlib `html.parser`.
    """
    # Lazy import — bs4 + lxml together are heavy.
    from bs4 import BeautifulSoup  # noqa: PLC0415

    try:
        soup = BeautifulSoup(content, "lxml")
    except Exception:
        soup = BeautifulSoup(content, "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    text = soup.get_text(separator="\n")
    # Collapse runs of blank lines so chunkers don't waste budget on whitespace.
    text = re.sub(r"\n\s*\n+", "\n\n", text)
    return text.strip()


def _parse_docx(content: bytes) -> str:
    """Extract paragraph text from a DOCX file using `python-docx`.
    Tables are flattened cell-by-cell with tab separators; embedded
    images are dropped. Encrypted documents raise ParseFailureError."""
    # Lazy import — python-docx (and its lxml dep) is several MB.
    import docx  # noqa: PLC0415
    from docx.opc.exceptions import PackageNotFoundError  # noqa: PLC0415

    try:
        document = docx.Document(io.BytesIO(content))
    except PackageNotFoundError as exc:
        raise ParseFailureError(f"invalid DOCX: {exc}") from exc

    parts: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append("\t".join(cells))
    return "\n\n".join(parts).strip()


_REGISTRY: dict[str, Callable[[bytes], str]] = {
    "text/plain": _parse_text,
    "text/markdown": _parse_markdown,
    "text/html": _parse_html,
    "application/pdf": _parse_pdf,
    _DOCX_MIME: _parse_docx,
}


def register_parser(mime_type: str, parser: Callable[[bytes], str]) -> None:
    """Operator hook to plug additional parsers at startup. Used to add
    OCR (Q-400-002) or specialised formats (xlsx, pptx) when they ship."""
    _REGISTRY[mime_type] = parser
